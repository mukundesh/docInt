from pathlib import Path

from more_itertools import first

from docint.para import Para

from ..region import Region
from ..shape import Box, Coord, Edge, Poly
from ..vision import Vision


@Vision.factory(
    "docx_generator",
    default_config={"stub": "docx", "ignore_page_idxs": [], "document_root": "."},
)
class DocxGenerator:
    def __init__(self, stub, ignore_page_idxs, document_root):
        self.stub = stub
        self.ignore_page_idxs = ignore_page_idxs
        self.document_root = Path(document_root)

    def find_paras(self, page):
        def is_center_aligned(line):
            padding = 0.1
            return (l_xmin + padding) < line.xmin < line.xmax < (l_xmax - padding)

        l_xmin = min((w.xmin for w in page.words), default=0.0)
        l_xmax = max((w.xmax for w in page.words), default=1.0)

        paras, para_lines = [], []

        for line_idx, line in enumerate(page.lines):
            print(f"{line_idx}:", end=" ")
            last_xmin = para_lines[-1].xmin if para_lines else None
            if (not line) or (not line.arranged_text().strip()):
                if para_lines:
                    print("build_para", end=" ")
                    paras.append(para_lines[:])
                    para_lines.clear()
            elif is_center_aligned(line) and len(line.words) < 3:
                print(f"center_aligned >{line.arranged_text()[:15]}...<", end=" ")
                para_lines.append(line)
                paras.append(para_lines[:])
                para_lines.clear()
            elif last_xmin is not None and abs(line.xmin - last_xmin) > 0.05:
                print(f"indented >{line.arranged_text()[:15]}...<", end=" ")
                paras.append(para_lines[:])
                para_lines.clear()
                para_lines.append(line)
            else:
                print(f"adding to Para[{len(paras)}] >{line.arranged_text()}<", end=" ")
                para_lines.append(line)
            print()

        if para_lines:
            print(f"creating para: of # lines: {len(para_lines)}")
            paras.append(para_lines[:])
            para_lines.clear()

        return paras

    def generate_page(self, document, page):
        def is_center_aligned(line):
            padding = 0.1
            return (l_xmin + padding) < line.xmin < line.xmax < (l_xmax - padding)

        l_xmin = min((w.xmin for w in page.words), default=0.0)
        l_xmax = max((w.xmax for w in page.words), default=1.0)

        from docx.shared import Inches
        from docx.text.paragraph import Paragraph
        from docx.text.run import Run

        def write_heading(document, line):
            line_text = line.arranged_text().strip()
            if line_text:
                document.add_heading(line.arranged_text().strip(), level=2)

        def write_para(document, para_lines, indent=False):
            para_text = "\n".join(ln.arranged_text().strip() for ln in para_lines)
            p = document.add_paragraph(para_text)
            if indent:
                print("INDENTING")
                p.paragraph_format.left_indent = Inches(0.5)

        paras = self.find_paras(page)
        prev_para_lines = None
        for para_lines in paras:
            if len(para_lines) == 1 and len(para_lines[0].words) < 3:
                write_heading(document, para_lines[0])
                prev_para_lines = None
            else:
                if not prev_para_lines:
                    indent = False
                else:
                    indent = (
                        True if (para_lines[0].xmin - prev_para_lines[0].xmin) > 0.05 else False
                    )
                write_para(document, para_lines, indent)
                prev_para_lines = para_lines
        return

    def __call__(self, doc):
        from docx import Document

        document = Document()
        for page_idx, page in enumerate(doc.pages):
            if page_idx in self.ignore_page_idxs:
                continue

            self.generate_page(document, page)
            document.add_page_break()

        document_path = self.document_root / f"{doc.pdf_name}.docx"
        document.save(document_path)
        return doc
